import docx
import pandas as pd
import csv
import os
from docx.oxml.ns import qn

# --- CONFIGURATION ---
input_file = r"C:\Users\hussy\Desktop\U2_Hamish(Tabs).docx"
output_csv = r"C:\Files\U2_Hamish\data.csv"
image_folder = r"C:\Files\U2_Hamish\images"

if not os.path.exists(image_folder):
    os.makedirs(image_folder)

def extract_to_final_csv(word_path, csv_path):
    doc = docx.Document(word_path)
    all_data = []
    img_counter = 1

    print("Extracting data into 7 columns (Page, Line, Word, Information, media, links, unzur)...")

    for table in doc.tables:
        for row in table.rows:
            try:
                # Ensure the row has 7 columns
                if len(row.cells) < 7:
                    continue

                # Mapping based on your 7-column request
                page_num    = row.cells[0].text.strip()
                line_num    = row.cells[1].text.strip()
                word_val    = row.cells[2].text.strip()
                info_cell   = row.cells[3]
                # Column 5 (index 4) handles local media if you put it there, 
                # but we will auto-fill it from extracted images.
                drive_links = row.cells[5].text.strip() # Column 6
                unzur_val   = row.cells[6].text.strip() # Column 7

                processed_lines = []
                row_media_files = []
                
                # Handle images and superscripts in the Information cell (Column 4)
                for para in info_cell.paragraphs:
                    line_html = ""
                    
                    for run in para.runs:
                        # 1. Inline Image Extraction
                        if 'pic:pic' in run.element.xml:
                            try:
                                for img in run.element.xpath('.//pic:pic'):
                                    rId = img.xpath('.//a:blip/@r:embed')[0]
                                    image_part = doc.part.related_parts[rId]
                                    
                                    img_name = f"img_{img_counter}.png"
                                    img_path = os.path.join(image_folder, img_name)
                                    
                                    with open(img_path, "wb") as f:
                                        f.write(image_part.blob)
                                    
                                    row_media_files.append(img_name)
                                    line_html += f'<img src="images/{img_name}" class="img-fluid"><br>'
                                    img_counter += 1
                            except:
                                continue

                        # 2. Text and Superscript Formatting
                        run_text = run.text
                        if run_text:
                            # Preserve spaces and handle superscripts
                            run_text = run_text.replace("  ", "&nbsp;&nbsp;")
                            if run.font.superscript:
                                run_text = f"<sup>{run_text}</sup>"
                            # Preserve bold formatting
                            if run.font.bold:
                                run_text = f"<b>{run_text}</b>"
                            # Preserve underline formatting
                            if run.font.underline:
                                run_text = f"<u>{run_text}</u>"
                            line_html += run_text
                    
                    # Apply centering if asterisk is present
                    if "*" in para.text:
                        line_html = f'<div style="text-align:center;">{line_html}</div>'
                    
                    if line_html:
                        processed_lines.append(line_html)

                full_info_html = "<br>".join(processed_lines)

                # --- ASSEMBLE ROW ---
                all_data.append({
                    'Page': page_num,
                    'Line': line_num,
                    'Word': word_val,
                    'Information': full_info_html,
                    'media': ", ".join(row_media_files), # Extracted local images
                    'links': drive_links,                # Drive links from Column 6
                    'unzur': unzur_val                  # Value from Column 7
                })

            except Exception as e:
                print(f"Error processing row: {e}")
                continue

    # Create DataFrame with exact column order
    df = pd.DataFrame(all_data, columns=['Page', 'Line', 'Word', 'Information', 'media', 'links', 'unzur'])
    
    # Save with UTF-8-SIG for proper Excel/Special Character support
    df.to_csv(csv_path, index=False, encoding='utf-8-sig', quoting=csv.QUOTE_ALL)
    print(f"Done! Created {csv_path}")

# Run
extract_to_final_csv(input_file, output_csv)